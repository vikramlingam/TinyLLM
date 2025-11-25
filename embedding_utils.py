import os
import pickle
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import pypdf
import docx
import io
from typing import List, Dict, Tuple

class RAGHandler:
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2', index_path: str = './rag_index'):
        self.model_name = model_name
        self.index_path = index_path
        self.model = None
        self.index = None
        self.chunks = [] # List to store text chunks corresponding to index vectors

    def load_model(self):
        """Loads the SentenceTransformer model."""
        if self.model is None:
            print(f"Loading embedding model: {self.model_name}...")
            self.model = SentenceTransformer(self.model_name)
            print("Embedding model loaded.")

    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Splits text into chunks with overlap."""
        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start += chunk_size - overlap
        
        return chunks

    def create_index(self, content, filename: str):
        """
        Chunks the provided content (text or bytes), computes embeddings, and builds the FAISS index in-memory.
        """
        self.load_model()
        
        text = ""
        if filename.lower().endswith('.pdf'):
            if isinstance(content, str):
                 # If path is passed (legacy support or if needed)
                 reader = pypdf.PdfReader(content)
            else:
                 # BytesIO
                 reader = pypdf.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif filename.lower().endswith('.docx'):
            if isinstance(content, str):
                doc = docx.Document(content)
            else:
                doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            # Assume text
            if isinstance(content, bytes):
                text = content.decode('utf-8')
            else:
                text = content
        
        print("Chunking text...")
        self.chunks = self._chunk_text(text)
        
        if not self.chunks:
            print("No text to index.")
            return

        print(f"Generated {len(self.chunks)} chunks. Computing embeddings...")
        embeddings = self.model.encode(self.chunks)
        
        # Initialize FAISS index
        dimension = embeddings.shape[1]
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(embeddings.astype('float32'))
        
        print(f"Index built with {self.index.ntotal} vectors.")

    def reset_index(self):
        """Clears the current index and chunks."""
        self.index = None
        self.chunks = []
        print("Index reset.")

    def query(self, query_text: str, k: int = 3) -> List[str]:
        """Retrieves the top-k most relevant chunks for the query."""
        if self.index is None or not self.chunks:
            return []
        
        # Ensure model is loaded for encoding the query
        self.load_model()
        
        query_vector = self.model.encode([query_text])
        distances, indices = self.index.search(query_vector.astype('float32'), k)
        
        results = []
        for idx in indices[0]:
            if idx != -1 and idx < len(self.chunks):
                results.append(self.chunks[idx])
                
        return results
